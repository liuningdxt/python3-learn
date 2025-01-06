import os
import sys

import pandas as pd
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QFont, QDragEnterEvent, QDropEvent, QIntValidator
from PyQt5.QtWidgets import QApplication, QMainWindow, QTableWidget, QTableWidgetItem, QVBoxLayout, QWidget, \
    QLabel, QHBoxLayout, QLineEdit, QMessageBox, QPushButton
from docx import Document

pd.set_option('display.max_rows', None)
pd.set_option('display.max_columns', None)


# PyQt5 主窗口
class TableWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.df = None
        self.setWindowTitle("准考证生成")
        self.setGeometry(100, 100, 2000, 1200)

        # 创建一个 QTableWidget 控件，填充效益表计算数据
        self.table = QTableWidget(self)
        self.fill_table(self.table)

        # 标签用于显示拖放提示
        self.label = QLabel("拖动准考证个人信息表Excel文件到此窗口", self)
        self.label.setAlignment(Qt.AlignCenter)
        font = QFont()
        font.setBold(True)
        font.setPointSize(10)
        self.label.setFont(font)

        self.label1 = QLabel("输入序号（整数）", self)
        self.label1.setAlignment(Qt.AlignCenter)
        self.label1.setFont(font)

        # 创建文本输入框
        self.input_box = QLineEdit(self)

        # 设置整数验证器
        int_validator = QIntValidator()
        self.input_box.setValidator(int_validator)

        # 创建按钮
        button = QPushButton('生成准考证信息', self)
        button.clicked.connect(self.generate_admission)

        # 创建水平布局
        h_layout = QHBoxLayout()
        h_layout.addWidget(self.label)  # 添加拖放提示标签
        h_layout.addStretch(1)
        h_layout.addWidget(self.label1)  # 添加序号标签
        h_layout.addWidget(self.input_box)
        h_layout.addWidget(button)  # 添加按钮
        h_layout.addStretch(1)

        # 创建一个垂直布局（以便进一步添加其他控件）
        v_layout = QVBoxLayout()
        v_layout.addLayout(h_layout)  # 添加水平布局（包含按钮和标签）
        v_layout.addWidget(self.table)  # 添加表格到布局

        # 创建一个QWidget容器，并设置为主窗口的中央部件
        container = QWidget()
        container.setLayout(v_layout)
        self.setCentralWidget(container)

        # 禁用所有单元格的编辑功能
        self.table.setEditTriggers(QTableWidget.NoEditTriggers)

        # 启用拖放功能
        self.setAcceptDrops(True)

    # 事件：拖动进入窗口时
    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    # 事件：文件被放下时
    def dropEvent(self, event: QDropEvent):
        file_url = event.mimeData().urls()[0].toLocalFile()
        if file_url.endswith(".xlsx") | file_url.endswith(".xls"):
            self.parse_excel(file_url)
        else:
            self.label.setText("只支持 excel 文件格式")

    # 解析 CSV 文件并显示在表格中
    def parse_excel(self, file_path):
        try:
            if file_path:
                # 使用 pandas 读取 Excel 文件并加载为 DataFrame
                # 显示 DataFrame 的内容在 QTableWidget 中
                self.display_df_in_table(file_path)

            self.label.setText(f"成功加载文件: {file_path}")
        except Exception as e:
            self.label.setText(f"文件加载失败: {str(e)}")

    def fill_table(self, table_widget):
        for row in range(table_widget.rowCount()):
            for col in range(table_widget.columnCount()):
                item = QTableWidgetItem('')
                table_widget.setItem(row, col, item)

    def display_df_in_table(self, filePath):
        df = pd.read_excel(filePath)
        df[['序号', '身份证号']] = df[['序号', '身份证号']].astype(str)
        self.df = df
        # 清空现有的表格内容
        self.table.setRowCount(0)
        self.table.setColumnCount(0)

        # 设置表格行列数
        self.table.setRowCount(len(df))
        self.table.setColumnCount(len(df.columns))

        # 设置表头
        self.table.setHorizontalHeaderLabels(df.columns)
        font = QFont()
        font.setBold(True)

        # 填充表格数据
        for row_idx, row in df.iterrows():
            for col_idx, value in enumerate(row):
                self.table.setItem(row_idx, col_idx, QTableWidgetItem(str(value)))

        # 设置表格样式
        self.table.setStyleSheet("""
                                    QTableWidget {
                                        border: 1px solid black;  /* 设置表格外边框 */
                                    }
                                    QHeaderView::section {
                                        border: 1px solid black;  /* 设置表头边框 */
                                    }
                                    QTableWidget::item {
                                        border: 1px solid black;  /* 设置单元格边框的粗细 */
                                    }
                                    QTableWidget::item:selected {
                                        background-color: #87CEFA;  /* 设置选中项的背景色为浅蓝色 */
                                    }
                                    """
                                 )

        # 自动调整行高
        self.table.resizeRowsToContents()
        # 设置列自适应宽度
        self.table.resizeColumnsToContents()

    def generate_admission(self):
        try:
            input_id = self.input_box.text()
            if not input_id:
                QMessageBox.warning(self, '输入错误', '序号输入不能为空！')
                return
            if self.df is None:
                QMessageBox.warning(self, '输入错误', '未导入正确的准考证个人信息文件！')
                return
            df = self.df
            if df.empty:
                QMessageBox.warning(self, '输入错误', '未导入正确的准考证个人信息文件！')
                return
            df_result = df[df['序号'] == input_id]
            if df_result.empty:
                QMessageBox.warning(self, '输入错误', '未获取到相应序号的准考证个人信息！')
                return

            # 获取当前文件的路径
            if getattr(sys, 'frozen', False):
                # 获取 .exe 文件的路径
                exe_path = sys.argv[0]
                # 获取 .exe 文件所在的目录
                application_path = os.path.dirname(os.path.abspath(exe_path))
            else:
                # 程序未打包
                application_path = os.path.dirname(os.path.abspath(__file__))

            # Excel 文件路径
            template_path = os.path.join(application_path, '模板.docx')

            # template_path = '模板.docx'
            result = df_result.iloc[0]
            name = result['姓名']
            sex = result['性别']
            id = result['序号']
            post = result['应聘岗位']
            card = result['身份证号']
            w_address = result['笔试地点']
            i_address = result['面试地点']
            output_path = f'{name}.docx'  # 输出文件
            replacements = {
                '[姓名]': name,
                '[性别]': sex,
                '[面试地点]': i_address,
                '[笔试地点]': w_address,
                '[准考证]': id,
                '[岗位]': post,
                '[身份证]': card
            }

            self.replace_in_word_template(template_path, output_path, replacements)

            msg = QMessageBox(self)
            msg.setIcon(QMessageBox.Information)  # 设置图标为信息图标
            msg.setText(f'已生成{name}的准考证信息')
            msg.setWindowTitle("信息")
            msg.setStandardButtons(QMessageBox.Ok)  # 设置按钮为“确定”
            msg.exec_()

        except Exception as e:
            QMessageBox.warning(self, '出错', f'生成准考证信息时出错：{str(e)}')

    def replace_in_word_template(self, template_path, output_path, replacements):
        """
        替换 Word 模板中的占位符标签
        :param template_path: 模板文件路径
        :param output_path: 输出文件路径
        :param replacements: 字典类型，包含标签和替换值
        """
        doc = Document(template_path)

        for placeholder, value in replacements.items():
            self.replace_placeholder(doc, placeholder, value)

        doc.save(output_path)

    def replace_placeholder(self, doc, placeholder, value):
        """
        替换文档中的占位符
        :param doc: word文档对象
        :param placeholder: 需要替换的标签（占位符）
        :param value: 替换的值
        """
        for para in doc.paragraphs:
            if placeholder in para.text:
                inline_shapes = para.runs
                for run in inline_shapes:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        inline_shapes = cell.paragraphs
                        for para in inline_shapes:
                            for run in para.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)


# 运行应用程序
app = QApplication(sys.argv)
window = TableWindow()
window.show()
sys.exit(app.exec_())
